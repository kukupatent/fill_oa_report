"""
fill_oa_report.py
=================
의견제출통지서(PDF) + OA 검토보고서 템플릿(DOCX) →
서지사항 표를 자동으로 채운 DOCX를 출력합니다.
출력 파일명은 자동 생성됩니다: [파이특허][관리번호] 1차OA 검토보고서.docx

사용법:
    python fill_oa_report.py <의견제출통지서.pdf> [출원서.rtf|.docx] <템플릿.docx>

예시 (의견제출통지서만):
    python fill_oa_report.py 의견제출통지서.pdf template.docx

예시 (출원서 포함):
    python fill_oa_report.py 의견제출통지서.pdf 출원서.rtf template.docx
    python fill_oa_report.py 의견제출통지서.pdf 출원서.docx template.docx

필요 라이브러리:
    pip install python-docx pdfplumber
"""

import sys, re, shutil, calendar, copy
from pathlib import Path
from datetime import date
import pdfplumber
from docx.oxml import OxmlElement
from docx import Document

# ── 상수 ─────────────────────────────────────
FONT_NAME = "SpoqaHanSans-Light"
FONT_SIZE_PT = 10          # 포인트 단위 (OOXML sz = half-point, 즉 ×2)

REJECTION_LABELS = [
    "신규성", "신규사항추가", "기재불비",
    "미완성 발명", "진보성", "기타(비법정 발명)"
]

# 특허법 조항 → 거절이유 매핑 (띄어쓰기 허용 패턴)
LAW_TO_REASON = {
    r"제\s*29\s*조\s*제\s*1\s*항": "신규성",
    r"제\s*29\s*조\s*제\s*2\s*항": "진보성",
    r"제\s*47\s*조\s*제\s*2\s*항": "신규사항추가",
    r"제\s*42\s*조":               "기재불비",
    r"미\s*완\s*성":               "미완성 발명",
    r"비\s*법\s*정\s*발\s*명":     "기타(비법정 발명)",
}

# 서지사항 표(표2) 레이아웃: {행: {고유셀_열: 필드키}}
# 고유셀 기준 열 인덱스 (레이블=짝수, 값=홀수)
TABLE_LAYOUT = {
    0: {1: "출원종류",       3: "출원국가"},
    1: {1: "출원번호",       3: "출원일"},
    2: {1: "당소관리번호",    3: "출원인"},
    3: {1: "발명자"},
    4: {1: "발명의 명칭"},
    5: {1: "통지서 발행일",   3: "OA 종류"},
    6: {1: "의견서 마감일",   3: "발명자 검토회신 요청일"},
    7: {1: "거절이유"},
}

LABEL_TEXTS = {
    "출원종류","출원국가","출원번호","출원일","당소관리번호","출원인",
    "발명자","발명의 명칭","통지서 발행일","OA 종류","의견서 마감일",
    "발명자 검토회신 요청일","거절이유",
}


# ── 1. PDF 파싱 ──────────────────────────────

def add_months(d: date, months: int) -> date:
    """날짜에 월 수를 더합니다 (월말 처리 포함)."""
    month = d.month - 1 + months
    year  = d.year + month // 12
    month = month % 12 + 1
    day   = min(d.day, calendar.monthrange(year, month)[1])
    return date(year, month, day)


def parse_rejection_from_table(pdf_path: str) -> tuple:
    """
    거절이유 표의 '관련 법조항' 및 '거절이유가 있는 부분' 컬럼에서 추출.
    모든 페이지를 순서대로 탐색하여 첫 번째 거절이유 표를 사용.

    반환값: (found_set, claim_map)
      - found_set: {"진보성", "기재불비", ...}
      - claim_map: {"진보성": "청구항 제1항 내지 제12항, 제20항", "기재불비": "청구항 제13항 내지 제19항", ...}
    """
    found = set()
    claim_map = {}  # reason → 청구항 텍스트
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for tbl in tables:
                    if not tbl or len(tbl[0]) < 2:
                        continue
                    header = [str(c or "") for c in tbl[0]]
                    if not any("법조항" in h for h in header):
                        continue
                    law_col   = next(i for i, h in enumerate(header) if "법조항" in h)
                    claim_col = next((i for i, h in enumerate(header) if "있는 부분" in h), None)
                    for row in tbl[1:]:
                        law_text   = str(row[law_col] or "")
                        claim_text = str(row[claim_col] or "").strip() if claim_col is not None else ""
                        # 줄바꿈 정리
                        claim_text = " ".join(claim_text.split())
                        for pattern, reason in LAW_TO_REASON.items():
                            if re.search(pattern, law_text):
                                found.add(reason)
                                if claim_text:
                                    claim_map[reason] = claim_text
                    if found:  # 거절이유 표 찾으면 중단
                        return found, claim_map
    except Exception:
        pass
    return found, claim_map


def parse_oa_pdf(pdf_path: str, template_path: str = "") -> dict:
    with pdfplumber.open(pdf_path) as pdf:
        full_text = "\n".join(p.extract_text() or "" for p in pdf.pages)

    def find(pattern, flags=0):
        m = re.search(pattern, full_text, flags)
        return m.group(1).strip() if m else ""

    # ── 기본 항목 (패턴 강화: 띄어쓰기·구분자 유연하게)
    info = {}

    # 출원번호
    info["출원번호"] = find(r"출\s*원\s*번\s*호\s+([\d\-]+)")

    # 출원일 (YYYY.MM.DD 또는 YYYY-MM-DD)
    raw = find(r"출\s*원\s*일\s*자?\s*[:\s]\s*(\d{4}[\.\-]\d{1,2}[\.\-]\d{1,2}\.?)")
    info["출원일"] = raw.rstrip(".").replace("-", ".")

    # 출원인 (스마트쿼트 제거)
    raw = find(r"출\s*원\s*인\s*성\s*명\s+(.+?)(?:\s*\(특허고객번호|$)", re.DOTALL)
    info["출원인"] = raw.split("\n")[0].strip().strip("\u2018\u2019\u201c\u201d'\"")

    # 발명자 (여러 명 모두 추출, 쉼표로 연결)
    names = re.findall(r"발\s*명\s*자\s*성\s*명\s+(\S+)", full_text)
    info["발명자"] = ", ".join(names)

    # 발명의 명칭
    raw = find(r"발\s*명\s*의\s*명\s*칭\s+(.+?)(?:\n발송번호|\n출원번호|\n1\.|$)", re.DOTALL)
    info["발명의 명칭"] = " ".join(raw.split())

    # 통지서 발행일 (발송일자)
    info["통지서 발행일"] = find(r"발송일자[:\s]+(\d{4}\.\d{2}\.\d{2}\.?)").rstrip(".")

    # 의견서 마감일 (제출기일)
    raw = find(r"제출기일[:\s]+(\d{4}\.\d{2}\.\d{2}\.?)")
    info["의견서 마감일"] = raw.rstrip(".") + "." if raw else ""

    # ── 당소관리번호: PDF 파일명의 [파이특허][관리번호] 패턴에서 추출
    # 예: _파이특허__DPA250072IPB_의견제출통지서.pdf
    #     [파이특허][DPA250072IPB] 의견제출통지서.pdf
    filename = Path(pdf_path).stem
    # 우선: [파이특허][XXX] 패턴
    m = re.search(r'\[파이특허\]\[([^\]]+)\]', filename)
    if not m:
        # fallback: 파일명 내 관리번호 패턴 직접 탐색 (DPA250072IPB 등 suffix 포함)
        m = re.search(r'([DIO]P[AMEC]\d+[A-Z]*)', filename, re.IGNORECASE)
    info["당소관리번호"] = m.group(1).upper() if m else ""

    # ── 발명자 검토회신 요청일: 통지서 발행일 + 2개월
    try:
        parts = info["통지서 발행일"].split(".")
        y, mo, d = int(parts[0]), int(parts[1]), int(parts[2])
        info["발명자 검토회신 요청일"] = add_months(date(y, mo, d), 2).strftime("%Y.%m.%d.")
    except Exception:
        info["발명자 검토회신 요청일"] = ""

    # ── OA 종류: 문서 제목으로 판별
    if re.search(r"거\s*절\s*결\s*정\s*서", full_text):
        info["OA 종류"] = "거절결정"
    else:
        info["OA 종류"] = "1차 OA"

    # ── 거절이유: 거절이유 표에서 직접 추출 (가장 정확)
    found_set, claim_map = parse_rejection_from_table(pdf_path)
    info["거절이유_set"] = found_set
    info["거절이유_청구항"] = claim_map  # {"진보성": "청구항 제1항 내지 ...", ...}

    # 표 추출 실패 시 본문 텍스트로 fallback
    if not info["거절이유_set"]:
        for pattern, reason in LAW_TO_REASON.items():
            if re.search(pattern, full_text):
                info["거절이유_set"].add(reason)

    return info


def _parse_rtf_to_text(rtf_path: str) -> str:
    """RTF 파일을 순수 텍스트로 변환 (한국어 특허 RTF 특화, cp949 인코딩)."""
    try:
        data = open(rtf_path, 'rb').read()
        text = data.decode('cp949')
    except Exception:
        return ""

    # 유니코드 이스케이프 \uNNNN 처리
    text = re.sub(
        r'\\u(-?\d+)\?',
        lambda m: chr(int(m.group(1))) if 0 <= int(m.group(1)) <= 0x10FFFF else '',
        text
    )
    # \r\n은 RTF 편집기 줄감기(word wrap) — 가장 먼저 공백 없이 제거
    text = re.sub(r'\r\n', '', text)
    # \par → 줄바꿈 (실제 단락 구분)
    text = re.sub(r'\\par\b', '\n', text)
    # RTF 서식 명령 제거 (\s? 제거 — catastrophic backtracking 방지)
    text = re.sub(r'\\[a-zA-Z]+\-?\d*', '', text)
    # 그룹 괄호 제거
    text = re.sub(r'[{}]', '', text)
    # 공백 정규화
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)

    # 청구범위 섹션만 추출한 뒤 줄 이어붙이기 (전체 4MB에 적용하면 느림)
    m = re.search(r'(【청구범위】.*?)(?:【요약서】|【요약】|$)', text, re.DOTALL)
    if m:
        section = m.group(1)
    else:
        section = text

    lines = section.split('\n')
    joined = []
    for line in lines:
        line = line.strip()
        if not line:
            joined.append('')
            continue
        if re.match(r'【.*?】', line):
            joined.append(line)
            continue
        if joined and joined[-1] and not re.search(r'[;,.]$', joined[-1]) \
                and not re.match(r'【.*?】', joined[-1]):
            # 앞 줄 끝 공백 제거 후 이어붙이기
            joined[-1] = joined[-1].rstrip() + line
        else:
            joined.append(line)

    result = '\n'.join(joined).strip()
    return result


def _parse_docx_to_text(docx_path: str) -> str:
    """DOCX 출원서에서 텍스트 추출 (단락 단위 줄바꿈)."""
    try:
        doc = Document(docx_path)
        return "\n".join(para.text for para in doc.paragraphs)
    except Exception:
        return ""


def parse_claims_from_application(app_path: str) -> tuple:
    """
    출원서 RTF 또는 DOCX에서 【청구범위】 섹션을 파싱합니다.
    반환: (claim1_text, claim1_with_header, all_claims_text)
      - claim1_text        : 청구항 1 본문 (헤더 제외)
      - claim1_with_header : 【청구항 1】 헤더 + 본문
      - all_claims_text    : 청구항 1~N 전체 (헤더 포함)
    """
    ext = Path(app_path).suffix.lower()

    if ext == '.rtf':
        full_text = _parse_rtf_to_text(app_path)
    elif ext == '.docx':
        full_text = _parse_docx_to_text(app_path)
    else:
        return "", "", ""

    if not full_text:
        return "", "", ""

    m = re.search(r'【청구범위】(.*?)(?:【요약서】|【요약】|$)', full_text, re.DOTALL)
    if not m:
        return "", "", ""

    raw = m.group(1).strip()

    # 청구항별 분리
    parts = re.split(r'(【청구항\s*\d+】)', raw)
    claims = {}
    for i in range(1, len(parts), 2):
        header = parts[i].strip()
        body   = parts[i + 1].strip() if i + 1 < len(parts) else ""
        body   = re.sub(r'\n{2,}', '\n', body).strip()
        num_m  = re.search(r'(\d+)', header)
        if num_m:
            claims[int(num_m.group(1))] = (header, body)

    if not claims:
        return "", "", ""

    claim1_text        = claims[1][1] if 1 in claims else ""
    claim1_with_header = f"{claims[1][0]}\n{claims[1][1]}" if 1 in claims else ""

    all_parts = []
    for num in sorted(claims.keys()):
        header, body = claims[num]
        all_parts.append(f"{header}\n{body}")
    all_claims_text = "\n\n".join(all_parts)

    return claim1_text, claim1_with_header, all_claims_text


def unique_cell(row, col_index: int):
    """병합 고려: col_index번째 고유 셀 반환"""
    seen, col = set(), 0
    for cell in row.cells:
        cid = id(cell._tc)
        if cid not in seen:
            seen.add(cid)
            if col == col_index:
                return cell
            col += 1
    return None


def apply_font(run, bold=False, underline=False, color_rgb=None):
    """run에 표준 폰트를 적용합니다. 추가 서식 옵션 지원."""
    from docx.oxml.ns import qn
    from lxml import etree
    rpr = run._r.find(qn('w:rPr'))
    if rpr is None:
        rpr = etree.SubElement(run._r, qn('w:rPr'))
        run._r.insert(0, rpr)
    # 폰트
    fonts = rpr.find(qn('w:rFonts'))
    if fonts is None:
        fonts = etree.SubElement(rpr, qn('w:rFonts'))
    fonts.set(qn('w:ascii'),    FONT_NAME)
    fonts.set(qn('w:eastAsia'), FONT_NAME)
    fonts.set(qn('w:hAnsi'),    FONT_NAME)
    fonts.set(qn('w:cs'),       FONT_NAME)
    # 크기 (half-point 단위)
    half_pt = str(FONT_SIZE_PT * 2)
    for tag in (qn('w:sz'), qn('w:szCs')):
        el = rpr.find(tag)
        if el is None:
            el = etree.SubElement(rpr, tag)
        el.set(qn('w:val'), half_pt)
    # 볼드
    b_el = rpr.find(qn('w:b'))
    if bold:
        if b_el is None:
            b_el = etree.SubElement(rpr, qn('w:b'))
    else:
        if b_el is not None:
            rpr.remove(b_el)
    # 밑줄
    u_el = rpr.find(qn('w:u'))
    if underline:
        if u_el is None:
            u_el = etree.SubElement(rpr, qn('w:u'))
        u_el.set(qn('w:val'), 'single')
    else:
        if u_el is not None:
            rpr.remove(u_el)
    # 글자색
    color_el = rpr.find(qn('w:color'))
    if color_rgb:
        if color_el is None:
            color_el = etree.SubElement(rpr, qn('w:color'))
        color_el.set(qn('w:val'), color_rgb)
    else:
        if color_el is not None:
            rpr.remove(color_el)


def set_cell_text(cell, text: str, bold=False, underline=False, color_rgb=None):
    """첫 번째 단락·run에 텍스트 설정, 나머지 run 제거, 서식 보존"""
    if not cell.paragraphs:
        return
    for para in cell.paragraphs[1:]:
        para._p.getparent().remove(para._p)
    para = cell.paragraphs[0]
    if not para.runs:
        run = para.add_run(text)
        apply_font(run, bold=bold, underline=underline, color_rgb=color_rgb)
        return
    para.runs[0].text = text
    apply_font(para.runs[0], bold=bold, underline=underline, color_rgb=color_rgb)
    for run in para.runs[1:]:
        run._r.getparent().remove(run._r)


def fill_rejection(cell, found_reasons: set):
    """
    각 단락의 run들을 하나로 합친 뒤, ◆/◇ 기호를 교체하고
    단락 전체를 단일 run으로 재작성합니다.
    공백·들여쓰기는 보존합니다.
    """
    for para in cell.paragraphs:
        if not para.runs:
            continue

        # 1) 단락 전체 텍스트 합치기
        full_text = "".join(run.text for run in para.runs)

        # 2) ◆/◇ 교체
        for label in REJECTION_LABELS:
            mark = "◆" if label in found_reasons else "◇"
            full_text = full_text.replace(f"◆{label}", f"{mark}{label}")
            full_text = full_text.replace(f"◇{label}", f"{mark}{label}")

        # 3) 첫 번째 run에 전체 텍스트 쓰고, 나머지 run 제거
        first_run = para.runs[0]
        apply_font(first_run)
        first_run.text = full_text
        for run in para.runs[1:]:
            run._r.getparent().remove(run._r)


# ── 3. 대응방안 요약 표 ─────────────────────

# 거절이유 표시 순서 및 그룹 정의
# 신규성+진보성은 함께 파싱되면 하나의 행으로 합침
REASON_GROUPS = [
    ({"신규성", "진보성"}, "신규성/진보성"),   # 둘 다 있으면 합침
    ({"신규성"},           "신규성"),
    ({"진보성"},           "진보성"),
    ({"기재불비"},         "기재불비"),
    ({"신규사항추가"},     "신규사항추가"),
    ({"미완성 발명"},      "미완성 발명"),
    ({"기타(비법정 발명)"},"기타(비법정 발명)"),
]

DEFAULT_FEASIBILITY = "다소 높음\n(60-65%)"

# 극복가능성 레이블 → (셀 텍스트, 배경색 hex)
FEASIBILITY_OPTIONS = {
    "매우 높음": ("매우 높음\n(80%)",    "5AFFAF"),
    "다소 높음": ("다소 높음\n(60-65%)", "91FFCA"),
    "중간":      ("중간\n(55-60%)",      "FFD899"),
}
# 기본값 레이블
DEFAULT_FEASIBILITY_LABEL = "다소 높음"


def _get_cell(row, col_idx):
    """고유셀 기준 col_idx번째 셀 반환."""
    seen, col = set(), 0
    for cell in row.cells:
        cid = id(cell._tc)
        if cid not in seen:
            seen.add(cid)
            if col == col_idx:
                return cell
            col += 1
    return None


def _set_response_cell(cell, text: str, justify: bool = False, indent_twips: int = 0):
    """
    대응방안 표 셀에 텍스트 설정. 기존 run 서식 복제. \n은 줄바꿈으로 처리.
    justify=True  → 양쪽맞춤 (w:jc both)
    indent_twips  → 첫줄 들여쓰기. 단, 【청구항 N】 헤더 줄은 들여쓰기 제외.
                    헤더가 있으면 헤더 단락 + 본문 단락 2개로 분리.
    """
    from docx.oxml.ns import qn as _qn

    # 1) 모든 단락에서 run 제거, 첫 단락 서식 템플릿 추출
    rpr_template = None
    for p_idx, para in enumerate(cell.paragraphs):
        for run in para.runs:
            if p_idx == 0 and rpr_template is None:
                rpr_el = run._r.find(_qn('w:rPr'))
                if rpr_el is not None:
                    rpr_template = copy.deepcopy(rpr_el)
            run._r.getparent().remove(run._r)

    # 2) 첫 단락 이후 단락 제거
    for para in cell.paragraphs[1:]:
        para._p.getparent().remove(para._p)

    # 3) 헤더/본문 분리
    #    【청구항 N】로 시작하면 첫 줄=헤더, 나머지=본문 (별도 단락)
    lines = text.split("\n")
    header_line = None
    body_lines  = lines
    if lines and re.match(r'【청구항\s*\d+】', lines[0].strip()):
        header_line = lines[0]
        body_lines  = lines[1:]

    def _make_rPr() -> OxmlElement:
        """run 속성: rpr_template 기반으로 FONT_NAME/SIZE 강제 지정."""
        rPr = copy.deepcopy(rpr_template) if rpr_template is not None else OxmlElement('w:rPr')
        # 기존 rFonts 제거 후 새로 지정 (테마 폰트 등 덮어쓰기)
        for old in rPr.findall(_qn('w:rFonts')):
            rPr.remove(old)
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(_qn('w:ascii'),    FONT_NAME)
        rFonts.set(_qn('w:eastAsia'), FONT_NAME)
        rFonts.set(_qn('w:hAnsi'),    FONT_NAME)
        rPr.insert(0, rFonts)
        # 크기 지정
        for old in rPr.findall(_qn('w:sz')):
            rPr.remove(old)
        for old in rPr.findall(_qn('w:szCs')):
            rPr.remove(old)
        sz = OxmlElement('w:sz')
        sz.set(_qn('w:val'), str(FONT_SIZE_PT * 2))
        szCs = OxmlElement('w:szCs')
        szCs.set(_qn('w:val'), str(FONT_SIZE_PT * 2))
        rPr.append(sz); rPr.append(szCs)
        return rPr

    def _write_lines_to_para(para, lns, apply_indent: bool):
        """단락 pPr 설정 후 lns를 w:br 연결로 씀."""
        pPr = para._p.find(_qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            para._p.insert(0, pPr)
        if justify:
            jc = pPr.find(_qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                pPr.append(jc)
            jc.set(_qn('w:val'), 'both')
        if apply_indent and indent_twips:
            ind = pPr.find(_qn('w:ind'))
            if ind is None:
                ind = OxmlElement('w:ind')
                pPr.append(ind)
            ind.set(_qn('w:firstLine'), str(indent_twips))

        for i, line in enumerate(lns):
            r_el = OxmlElement('w:r')
            r_el.append(_make_rPr())
            t_el = OxmlElement('w:t')
            t_el.text = line
            t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r_el.append(t_el)
            para._p.append(r_el)
            if i < len(lns) - 1:
                br_r = OxmlElement('w:r')
                br_el = OxmlElement('w:br')
                br_r.append(br_el)
                para._p.append(br_r)

    first_para = cell.paragraphs[0]

    if header_line is not None:
        # 헤더 단락: 들여쓰기 없음
        _write_lines_to_para(first_para, [header_line], apply_indent=False)
        # 본문: 각 줄을 별도 단락으로 (모두 들여쓰기 적용)
        insert_after = first_para._p
        for line in body_lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue
            new_p = OxmlElement('w:p')
            insert_after.addnext(new_p)
            insert_after = new_p
            pPr = OxmlElement('w:pPr')
            if justify:
                jc = OxmlElement('w:jc')
                jc.set(_qn('w:val'), 'both')
                pPr.append(jc)
            if indent_twips:
                ind = OxmlElement('w:ind')
                ind.set(_qn('w:firstLine'), str(indent_twips))
                pPr.append(ind)
            new_p.append(pPr)
            r_el = OxmlElement('w:r')
            r_el.append(_make_rPr())
            t_el = OxmlElement('w:t')
            t_el.text = line_stripped
            t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r_el.append(t_el)
            new_p.append(r_el)
    else:
        # 헤더 없음: 전체를 하나의 단락에
        _write_lines_to_para(first_para, body_lines, apply_indent=True)


def fill_response_table(doc, found_reasons: set, override_feasibility: dict = None):
    """
    대응방안 요약 표(표3)를 파싱된 거절이유로 채웁니다.
    override_feasibility: {거절이유 레이블: "매우 높음" | "다소 높음" | "중간"}
    """
    import copy
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    tbl = doc.tables[3]

    # 거절이유별 행 목록 계산
    remaining = set(found_reasons)
    rows_to_fill = []

    for reason_set, label in REASON_GROUPS:
        if reason_set <= remaining:
            f_label = (override_feasibility or {}).get(label, DEFAULT_FEASIBILITY_LABEL)
            rows_to_fill.append((label, f_label))
            remaining -= reason_set

    for reason in sorted(remaining):
        f_label = (override_feasibility or {}).get(reason, DEFAULT_FEASIBILITY_LABEL)
        rows_to_fill.append((reason, f_label))

    if not rows_to_fill:
        return

    template_tr = copy.deepcopy(tbl.rows[1]._tr)

    for row in tbl.rows[1:]:
        tbl._tbl.remove(row._tr)

    for label, f_label in rows_to_fill:
        new_tr = copy.deepcopy(template_tr)
        tbl._tbl.append(new_tr)
        new_row = tbl.rows[-1]

        # 열0: 거절이유
        c0 = _get_cell(new_row, 0)
        if c0:
            _set_response_cell(c0, label)

        # 열1: 대응방안 요약 (빈칸)
        c1 = _get_cell(new_row, 1)
        if c1:
            _set_response_cell(c1, "")

        # 열2: 극복가능성 (텍스트 + 배경색)
        c2 = _get_cell(new_row, 2)
        if c2:
            text, bg_color = FEASIBILITY_OPTIONS.get(
                f_label, FEASIBILITY_OPTIONS[DEFAULT_FEASIBILITY_LABEL]
            )
            _set_response_cell(c2, text)
            # 배경색 적용
            tcPr = c2._tc.find(_qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                c2._tc.insert(0, tcPr)
            shd = tcPr.find(_qn('w:shd'))
            if shd is None:
                shd = OxmlElement('w:shd')
                tcPr.append(shd)
            shd.set(_qn('w:val'),   'clear')
            shd.set(_qn('w:color'), 'auto')
            shd.set(_qn('w:fill'),  bg_color)


# ── 3-b. OA 내용 분석 표 ─────────────────────

# 표4 구조:
#   행0: 헤더 (고정)
#   행1: 현재 거절 이유 행1 (거절이유 첫 번째)
#   행2: 현재 거절 이유 행2 (거절이유 두 번째, 없으면 비움)
#   행3: 인용발명 행 (고정)
#
# 거절이유별 "거절 이유가 있는 부분" 컬럼 기본값 (빈칸으로 두어도 됨)
OA_ANALYSIS_REASON_LABEL = {
    "신규성":           "신규성",
    "진보성":           "진보성",
    "신규성/진보성":     "신규성/진보성",
    "기재불비":         "기재불비",
    "신규사항추가":      "신규사항추가",
    "미완성 발명":       "미완성 발명",
    "기타(비법정 발명)": "기타(비법정 발명)",
}


def fill_oa_analysis_table(doc, info: dict):
    """
    OA 내용 분석 표(표4)를 파싱된 거절이유로 채웁니다.
    - col1: 거절이유가 있는 부분 (청구항)
    - col2: 거절이유
    - 거절이유 1개: 행1만 사용, 행2 삭제 + vMerge 해제
    - 거절이유 2개: 행1, 행2 각각 기재
    """
    tbl = doc.tables[4]
    found_reasons = info.get("거절이유_set", set())
    claim_map     = info.get("거절이유_청구항", {})

    # REASON_GROUPS 순서로 레이블 목록 생성
    remaining = set(found_reasons)
    reason_labels = []
    for reason_set, label in REASON_GROUPS:
        if reason_set <= remaining:
            # 신규성/진보성 합친 경우: 두 청구항 텍스트 합치기
            claims = [claim_map.get(r, "") for r in reason_set if claim_map.get(r)]
            claim_map[label] = "\n".join(claims) if claims else ""
            reason_labels.append(label)
            remaining -= reason_set
    for r in sorted(remaining):
        reason_labels.append(r)

    def get_unique_cells(row):
        seen = set(); cells = []
        for cell in row.cells:
            cid = id(cell._tc)
            if cid not in seen:
                seen.add(cid); cells.append(cell)
        return cells

    from docx.oxml.ns import qn as _qn

    row1 = tbl.rows[1]
    row2 = tbl.rows[2]
    cells1 = get_unique_cells(row1)
    cells2 = get_unique_cells(row2)

    if len(reason_labels) <= 1:
        # ── 거절이유 1개: 행1만 사용, 행2 삭제
        label = reason_labels[0] if reason_labels else ""
        if len(cells1) > 1:
            _set_response_cell(cells1[1], claim_map.get(label, ""))
        if len(cells1) > 2:
            _set_response_cell(cells1[2], label)

        # 행1 col0 vMerge 해제 (단독 셀로)
        tcPr0 = cells1[0]._tc.find(_qn('w:tcPr'))
        if tcPr0 is not None:
            vMerge = tcPr0.find(_qn('w:vMerge'))
            if vMerge is not None:
                tcPr0.remove(vMerge)

        # 행2 삭제
        tbl._tbl.remove(row2._tr)

        # 기재불비만인 경우: 인용발명 행도 삭제 (신규성/진보성 없으면 인용발명 불필요)
        needs_citation = found_reasons - {"기재불비", "신규사항추가", "미완성 발명", "기타(비법정 발명)"}
        if not needs_citation:
            # 행2 삭제 후 인용발명 행은 현재 행2(0-indexed)
            if len(tbl.rows) > 2:
                tbl._tbl.remove(tbl.rows[2]._tr)

    else:
        # ── 거절이유 2개: 행1, 행2 각각 채우기
        for i, (row, cells) in enumerate([(row1, cells1), (row2, cells2)]):
            label = reason_labels[i] if i < len(reason_labels) else ""
            if len(cells) > 1:
                _set_response_cell(cells[1], claim_map.get(label, ""))
            if len(cells) > 2:
                _set_response_cell(cells[2], label)


# ── 3. DOCX 채우기 ───────────────────────────

def _insert_claims_after_para(target_para, all_claims_text: str):
    """
    target_para 바로 뒤에 청구항을 삽입합니다.
    구조:
      - 【청구항 N】 헤더: 들여쓰기 없음, 양쪽맞춤 (단락 1개)
      - 본문 각 줄: 모두 들여쓰기 1.41cm, 양쪽맞춤 (줄마다 별도 단락)
    """
    from docx.oxml.ns import qn as _qn

    INDENT = 799  # 1.41cm in twips
    insert_after = target_para._p

    parts = re.split(r'(【청구항\s*\d+】)', all_claims_text)
    blocks = []
    for i in range(1, len(parts), 2):
        header = parts[i].strip()
        body   = parts[i + 1].strip() if i + 1 < len(parts) else ""
        blocks.append((header, body))

    if not blocks:
        blocks = [("", all_claims_text)]

    def _make_para(line: str, apply_indent: bool) -> OxmlElement:
        """단일 줄 → 단락 1개."""
        new_p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(_qn('w:val'), 'both')
        pPr.append(jc)
        if apply_indent:
            ind = OxmlElement('w:ind')
            ind.set(_qn('w:firstLine'), str(INDENT))
            pPr.append(ind)
        new_p.append(pPr)
        if line:
            r_el = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(_qn('w:ascii'), FONT_NAME)
            rFonts.set(_qn('w:eastAsia'), FONT_NAME)
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(_qn('w:val'), str(FONT_SIZE_PT * 2))
            szCs = OxmlElement('w:szCs')
            szCs.set(_qn('w:val'), str(FONT_SIZE_PT * 2))
            rPr.append(sz); rPr.append(szCs)
            r_el.append(rPr)
            t_el = OxmlElement('w:t')
            t_el.text = line
            t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r_el.append(t_el)
            new_p.append(r_el)
        return new_p

    for header, body in blocks:
        # 헤더 단락: 들여쓰기 없음
        if header:
            p_header = _make_para(header, apply_indent=False)
            insert_after.addnext(p_header)
            insert_after = p_header
        # 본문: 각 줄을 별도 단락으로, 모두 들여쓰기 적용
        if body:
            for line in body.split('\n'):
                line = line.strip()
                if not line:
                    continue
                p_line = _make_para(line, apply_indent=True)
                insert_after.addnext(p_line)
                insert_after = p_line


def fill_current_claim_table(doc, claim1_with_header: str):
    """
    표5 '현재 청구항 (독립항 제 1항)' 아래 행(행1)에
    【청구항 1】 헤더 포함 본문을 양쪽맞춤 + 들여쓰기(1.41cm)로 채웁니다.
    """
    if not claim1_with_header or len(doc.tables) < 6:
        return
    tbl = doc.tables[5]
    if len(tbl.rows) < 2:
        return
    cell = tbl.rows[1].cells[0]
    _set_response_cell(cell, claim1_with_header, justify=True, indent_twips=799)


def fill_amendment_table(doc, claim1_with_header: str, all_claims_text: str):
    """
    표6 '보정안' 아래 행(행1)에 청구항 1만 양쪽맞춤 + 들여쓰기로 채웁니다.
    '[첨부 1] 당소 보정안' 섹션: 기존 청구항 단락 삭제 후 새 전체 청구범위 삽입.
    """
    # 표6: 청구항 1만
    if claim1_with_header and len(doc.tables) >= 7:
        tbl = doc.tables[6]
        if len(tbl.rows) >= 2:
            cell = tbl.rows[1].cells[0]
            _set_response_cell(cell, claim1_with_header, justify=True, indent_twips=799)

    # [첨부 1] 당소 보정안: 청구범위 삽입
    if not all_claims_text:
        return

    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == "【특허청구범위】":
            _insert_claims_after_para(para, all_claims_text)
            break



def fill_inventor_review_para(doc, info: dict):
    """
    '5. 발명자 검토 요청 사항' 단락의 날짜를 파싱된 값으로 교체합니다.

    원본 구조:
      run0  : "위의 내용을 검토하시어...바랍니다. "  (고정 텍스트)
      run1~9 : "상기 건의...년 " + 의견서 마감일(볼드+빨강) + "이므로 늦어도 "
      run11~21: 발명자 검토회신 요청일(볼드+파랑) + "까지..."
      run22~23: "위 내용에..." + "감사합니다. " (고정)

    → 날짜 부분 run을 각각 1개씩 새 run으로 교체
    """
    from docx.oxml.ns import qn as _qn
    import copy

    deadline    = info.get("의견서 마감일", "")       # "2026.01.23."
    review_date = info.get("발명자 검토회신 요청일", "")  # "2025.11.23."

    def _fmt_korean(d: str) -> str:
        """2026.01.23. → 2026년 1월 23일"""
        m = re.match(r'(\d{4})\.(\d{1,2})\.(\d{1,2})\.?', d.strip())
        return f"{m.group(1)}년 {int(m.group(2))}월 {int(m.group(3))}일" if m else d

    deadline_kr    = _fmt_korean(deadline)
    review_date_kr = _fmt_korean(review_date)

    # 대상 단락 탐색
    target_para = None
    for para in doc.paragraphs:
        if "의견서 제출기일" in para.text or "검토의견을 부탁드립니다" in para.text:
            target_para = para
            break
    if target_para is None:
        return

    runs = target_para.runs

    def _make_run(template_run, text: str) -> OxmlElement:
        """template_run의 서식을 복사해 text만 교체한 새 w:r 반환."""
        new_r = copy.deepcopy(template_run._r)
        t_el = new_r.find(_qn('w:t'))
        if t_el is None:
            t_el = OxmlElement('w:t')
            new_r.append(t_el)
        t_el.text = text
        t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        return new_r

    # ── 의견서 마감일 교체: run1~9를 run1 서식으로 2개(년 부분 + 볼드 날짜)로 교체
    # run1: "상기 건의 의견서 제출기일은 YYYY년 " (일반)
    # run6~9: "M월 DD일" (볼드+빨강)
    # 전략: run1에 "상기 건의 의견서 제출기일은 " 쓰고,
    #        run1 뒤에 년 run, 볼드 run 순으로 삽입 후 run2~9 삭제

    if len(runs) >= 10:
        # 년도 부분 파싱
        year_deadline    = deadline_kr.split("년")[0] + "년 "   # "2026년 "
        month_day_dl     = deadline_kr.split("년 ")[1]           # "1월 23일"

        # run1 → "상기 건의 의견서 제출기일은 YYYY년 "
        runs[1]._r.find(_qn('w:t')).text = "상기 건의 의견서 제출기일은 " + year_deadline

        # run6(볼드+빨강 서식) → "M월 DD일"
        runs[6]._r.find(_qn('w:t')).text = month_day_dl

        # run2~5, run7~9 삭제 (합쳐진 내용이 run1, run6에)
        for run in list(runs[2:6]) + list(runs[7:10]):
            run._r.getparent().remove(run._r)

        # run 목록 갱신 후 마감일 뒤 텍스트 run 확인
        runs = target_para.runs

    # ── 검토회신 요청일 교체: "이므로 늦어도 YYYY년 " + 볼드 "MM월 DD일"
    # 현재 runs에서 "이므로" run, 년도 run, 볼드 run 찾기
    if len(runs) >= 8:
        year_review   = review_date_kr.split("년")[0] + "년 "   # "2025년 "
        month_day_rv  = review_date_kr.split("년 ")[1]           # "11월 23일"

        # "이므로 늦어도 " run 탐색
        for i, run in enumerate(runs):
            if "이므로" in run.text:
                # 다음에 년도 run들 + 볼드 run이 있음
                # 이 run에 "이므로 늦어도 YYYY년 " 씀
                runs[i]._r.find(_qn('w:t')).text = "이므로 늦어도 " + year_review
                # i+1부터 볼드 run 찾기
                bold_start = None
                for j in range(i+1, len(runs)):
                    rPr = runs[j]._r.find(_qn('w:rPr'))
                    b   = rPr.find(_qn('w:b')) if rPr is not None else None
                    if b is not None:
                        bold_start = j
                        break
                if bold_start is not None:
                    # 중간 일반 run들(년도 조각들) 삭제
                    for run in list(runs[i+1:bold_start]):
                        run._r.getparent().remove(run._r)
                    runs = target_para.runs
                    # bold_start 재탐색
                    for j, run in enumerate(runs):
                        rPr = run._r.find(_qn('w:rPr'))
                        b   = rPr.find(_qn('w:b')) if rPr is not None else None
                        clr = rPr.find(_qn('w:color')) if rPr is not None else None
                        clr_val = clr.get(_qn('w:val')) if clr is not None else None
                        if b is not None and clr_val == "0070C0":
                            # 이 run → "MM월 DD일"
                            runs[j]._r.find(_qn('w:t')).text = month_day_rv
                            # 뒤따르는 같은 색 볼드 run들 삭제
                            k = j + 1
                            while k < len(runs):
                                rPr2 = runs[k]._r.find(_qn('w:rPr'))
                                b2   = rPr2.find(_qn('w:b')) if rPr2 is not None else None
                                clr2 = rPr2.find(_qn('w:color')) if rPr2 is not None else None
                                cv2  = clr2.get(_qn('w:val')) if clr2 is not None else None
                                if b2 is not None and cv2 == "0070C0":
                                    runs[k]._r.getparent().remove(runs[k]._r)
                                    runs = target_para.runs
                                else:
                                    break
                            break
                break


def fill_docx(template_path: str, info: dict, output_path: str,
              app_pdf_path: str = "", override_feasibility: dict = None):
    shutil.copy(template_path, output_path)
    doc = Document(output_path)

    if len(doc.tables) < 3:
        raise ValueError(f"서지사항 표를 찾을 수 없습니다. (표 개수: {len(doc.tables)})")

    # ── 표지 채우기 (표1, 0-indexed)
    cover = doc.tables[1]

    def set_cover_cell(row_idx, text):
        """표지 값 셀: 기존 run들을 합쳐 첫 run에 쓰고 서식 보존."""
        seen=set(); cells=[]
        for cell in cover.rows[row_idx].cells:
            cid=id(cell._tc)
            if cid not in seen:
                seen.add(cid); cells.append(cell)
        val_cell = cells[1]
        para = val_cell.paragraphs[0]
        if not para.runs:
            return
        para.runs[0].text = text
        for run in para.runs[1:]:
            run._r.getparent().remove(run._r)

    def set_cover_title(text):
        """표지 제목 셀: 기존 run들을 합쳐 첫 run에 쓰고 서식 보존."""
        para = cover.rows[0].cells[0].paragraphs[0]
        if not para.runs:
            return
        para.runs[0].text = text
        for run in para.runs[1:]:
            run._r.getparent().remove(run._r)

    # 제목
    oa_type = info.get("OA 종류", "")
    if oa_type == "거절결정":
        cover_title = "거절결정 검토 보고서"
    else:
        cover_title = "1차 OA 검토 보고서"
    set_cover_title(cover_title)

    # 의뢰인 → 출원인
    set_cover_cell(1, info.get("출원인", ""))

    # 작성일 → 오늘
    from datetime import date as _date
    today = _date.today().strftime("%Y. %m. %d.").replace(" 0", " ")
    set_cover_cell(2, today)

    # 담당자 → 빈칸
    set_cover_cell(3, "")

    tbl = doc.tables[2]
    fill_map = {
        "출원종류":             "특허",
        "출원국가":             "한국",
        "출원번호":             info.get("출원번호", ""),
        "출원일":              info.get("출원일", ""),
        "당소관리번호":          info.get("당소관리번호", ""),
        "출원인":              info.get("출원인", ""),
        "발명자":              info.get("발명자", ""),
        "발명의 명칭":          info.get("발명의 명칭", ""),
        "통지서 발행일":         info.get("통지서 발행일", ""),
        "OA 종류":             info.get("OA 종류", ""),
        "의견서 마감일":         info.get("의견서 마감일", ""),
        "발명자 검토회신 요청일": info.get("발명자 검토회신 요청일", ""),
    }
    found_reasons = info.get("거절이유_set", set())

    for row_idx, col_map in TABLE_LAYOUT.items():
        if row_idx >= len(tbl.rows):
            continue
        row = tbl.rows[row_idx]
        for col_idx, field_key in col_map.items():
            cell = unique_cell(row, col_idx)
            if cell is None or cell.text.strip() in LABEL_TEXTS:
                continue
            if field_key == "거절이유":
                fill_rejection(cell, found_reasons)
            elif field_key == "의견서 마감일":
                set_cell_text(cell, fill_map.get(field_key, ""),
                              bold=True, underline=True)
            elif field_key == "발명자 검토회신 요청일":
                set_cell_text(cell, fill_map.get(field_key, ""),
                              bold=True, underline=True, color_rgb="FF0000")
            else:
                set_cell_text(cell, fill_map.get(field_key, ""))

    # ── 대응방안 요약 표 채우기 (표3, 0-indexed)
    fill_response_table(doc, info.get("거절이유_set", set()), override_feasibility=override_feasibility)

    # ── OA 내용 분석 표 채우기 (표4, 0-indexed)
    fill_oa_analysis_table(doc, info)

    # ── 출원서 청구항 채우기 (표5, 표6, [첨부 1])
    if app_pdf_path:
        claim1_text, claim1_with_header, all_claims_text = parse_claims_from_application(app_pdf_path)
        fill_current_claim_table(doc, claim1_with_header)
        fill_amendment_table(doc, claim1_with_header, all_claims_text)

    # ── 5. 발명자 검토 요청 사항 단락 채우기
    fill_inventor_review_para(doc, info)

    doc.save(output_path)


# ── 4. 진입점 ────────────────────────────────

def make_output_filename(info: dict) -> str:
    """OA 종류와 관리번호로 출력 파일명 자동 생성."""
    mgmt_no = info.get("당소관리번호", "").upper()
    oa_type = info.get("OA 종류", "")

    if oa_type == "거절결정":
        suffix = "거절결정 검토보고서"
    else:
        suffix = "1차OA 검토보고서"

    return f"[파이특허][{mgmt_no}] {suffix}.docx"


def main():
    args = sys.argv[1:]
    if len(args) < 2:
        print(__doc__)
        sys.exit(1)

    # 인수 파싱:
    #   args[0]  = 의견제출통지서 PDF (필수)
    #   args[1]  = 출원서 RTF 또는 DOCX (선택) — .rtf면 무조건 출원서,
    #              .docx이고 args[2]도 있으면 출원서
    #   args[-1] = 템플릿 DOCX (필수)
    #
    # 출력 파일명은 항상 자동 생성 ([파이특허][관리번호] 1차OA 검토보고서.docx)

    pdf_path = args[0]  # 의견제출통지서

    app_pdf_path = ""
    if len(args) >= 2 and args[1].lower().endswith(".rtf"):
        app_pdf_path  = args[1]
        template_path = args[2] if len(args) >= 3 else None
    elif len(args) >= 3 and args[1].lower().endswith(".docx"):
        app_pdf_path  = args[1]
        template_path = args[2]
    else:
        template_path = args[1] if len(args) >= 2 else None

    if not template_path:
        print("❌ 템플릿 DOCX 경로를 지정해주세요.")
        print(__doc__)
        sys.exit(1)

    print(f"[1/3] 파싱: {pdf_path}")
    info = parse_oa_pdf(pdf_path, template_path)
    print("      결과:")
    for k, v in info.items():
        print(f"        {'거절이유 항목' if k == '거절이유_set' else k}: {v}")

    output_path = make_output_filename(info)

    print(f"\n[2/3] 템플릿: {template_path}")
    if app_pdf_path:
        print(f"      출원서: {app_pdf_path}")
    print(f"[3/3] 저장:   {output_path}")
    fill_docx(template_path, info, output_path, app_pdf_path=app_pdf_path)
    print(f"\n✓ 완료: {output_path}")


if __name__ == "__main__":
    main()
